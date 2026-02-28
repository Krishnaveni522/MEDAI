import sqlite3
from docx import Document
from pathlib import Path
from typing import Dict, List, Tuple, Optional
import os

DB_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'drug_information.db')

def init_database():
    """Initialize SQLite database with drug_information table"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS drug_information (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            drug TEXT NOT NULL UNIQUE,
            instructions TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    conn.commit()
    conn.close()

def load_from_docx_to_db(docx_path: str) -> int:
    """
    Load medication instructions from a .docx file into the database.
    Converts DOCX content (text + images) to HTML for storage.
    Returns the number of records inserted.
    """
    path = Path(docx_path)
    if not path.exists():
        return 0

    doc = Document(str(path))
    instructions = {}
    
    # Create images directory
    images_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'static', 'drug_images')
    os.makedirs(images_dir, exist_ok=True)

    # Read tables and convert content to HTML
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2 and cells[0]:
                name = cells[0]
                if name.lower() not in ['drug', 'medicine', 'name']:  # Skip header rows
                    # Convert instruction cell to HTML with embedded images
                    html_content = _convert_cell_to_html(row.cells[1], name, images_dir)
                    instructions[name] = html_content

    # Fallback: parse paragraphs
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    i = 0
    while i < len(paras):
        line = paras[i]
        # pattern "Name: Instruction"
        if ":" in line:
            key, val = line.split(":", 1)
            key, val = key.strip(), val.strip()
            if key and key.lower() not in ['drug', 'medicine', 'name']:
                if val:
                    instructions.setdefault(key, val)
                else:
                    # next paragraph as value
                    if i + 1 < len(paras):
                        instructions.setdefault(key, paras[i + 1])
                        i += 1
        # pattern "Name - Instruction" or tab-separated
        elif "-" in line:
            key, val = line.split("-", 1)
            key, val = key.strip(), val.strip()
            if key and key.lower() not in ['drug', 'medicine', 'name']:
                instructions.setdefault(key, val)
        elif "\t" in line:
            key, val = line.split("\t", 1)
            key, val = key.strip(), val.strip()
            if key and key.lower() not in ['drug', 'medicine', 'name']:
                instructions.setdefault(key, val)
        i += 1

    # Insert into database
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    count = 0
    for drug, html_instr in instructions.items():
        try:
            cursor.execute(
                'INSERT OR REPLACE INTO drug_information (drug, instructions) VALUES (?, ?)',
                (drug, html_instr)
            )
            count += 1
            print(f"Inserted: {drug}")
        except sqlite3.Error as e:
            print(f"Error inserting {drug}: {e}")
    
    conn.commit()
    conn.close()
    
    return count

def _convert_cell_to_html(cell, drug_name: str, images_dir: str) -> str:
    """Convert a DOCX table cell to HTML with embedded images"""
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn
    
    html_parts = []
    img_counter = 0
    
    for paragraph in cell.paragraphs:
        para_parts = []
        
        for run in paragraph.runs:
            # Check if run contains an image
            if 'graphicData' in run._element.xml or 'pic:pic' in run._element.xml:
                # Extract and save image from this specific run
                try:
                    # Get the image relationship ID from the run's drawing element
                    inline_shapes = run._element.xpath('.//a:blip')
                    if inline_shapes:
                        for inline in inline_shapes:
                            embed_id = inline.get(qn('r:embed'))
                            if embed_id:
                                # Get the image from the relationship
                                image_part = paragraph.part.related_parts[embed_id]
                                image_data = image_part.blob
                                
                                # Save image with unique filename
                                safe_name = drug_name.replace(' ', '_').replace('/', '_').replace('\\', '_')
                                img_filename = f"{safe_name}_{img_counter}.png"
                                img_path = os.path.join(images_dir, img_filename)
                                
                                with open(img_path, 'wb') as img_file:
                                    img_file.write(image_data)
                                
                                # Add image HTML tag
                                para_parts.append(f'<br><img src="/static/drug_images/{img_filename}" style="max-width: 400px; display: block; margin: 10px 0; border-radius: 5px;"><br>')
                                img_counter += 1
                                print(f"  Extracted image: {img_filename}")
                except Exception as e:
                    print(f"  Error extracting image: {e}")
            else:
                # Regular text
                text = run.text
                if text:
                    # Apply basic formatting
                    if run.bold:
                        text = f'<strong>{text}</strong>'
                    if run.italic:
                        text = f'<em>{text}</em>'
                    para_parts.append(text)
        
        if para_parts:
            html_parts.append(''.join(para_parts) + '<br>')
    
    return '\n'.join(html_parts)

def get_all_drugs() -> List[Tuple[int, str, str]]:
    """Get all drugs from the database"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    cursor.execute('SELECT id, drug, instructions FROM drug_information ORDER BY drug')
    results = cursor.fetchall()
    
    conn.close()
    return results

def get_all_drugs_dict() -> Dict[str, str]:
    """Get all drugs as a dictionary {drug_name: instructions}"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    cursor.execute('SELECT drug, instructions FROM drug_information ORDER BY drug')
    results = cursor.fetchall()
    
    conn.close()
    return {drug: instr for drug, instr in results}

def add_drug(drug: str, instructions: str) -> bool:
    """Add a new drug to the database"""
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        
        cursor.execute(
            'INSERT INTO drug_information (drug, instructions) VALUES (?, ?)',
            (drug.strip(), instructions.strip())
        )
        
        conn.commit()
        conn.close()
        return True
    except sqlite3.IntegrityError:
        return False  # Drug already exists
    except Exception as e:
        print(f"Error adding drug: {e}")
        return False

def update_drug(drug_id: int, drug: str, instructions: str) -> bool:
    """Update an existing drug in the database"""
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        
        cursor.execute(
            'UPDATE drug_information SET drug = ?, instructions = ? WHERE id = ?',
            (drug.strip(), instructions.strip(), drug_id)
        )
        
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        print(f"Error updating drug: {e}")
        return False

def delete_drug(drug_id: int) -> bool:
    """Delete a drug from the database"""
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        
        cursor.execute('DELETE FROM drug_information WHERE id = ?', (drug_id,))
        
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        print(f"Error deleting drug: {e}")
        return False

def search_drug(query: str) -> Optional[Tuple[str, str]]:
    """Search for a drug by name (exact match)"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    cursor.execute(
        'SELECT drug, instructions FROM drug_information WHERE LOWER(drug) = LOWER(?)',
        (query.strip(),)
    )
    result = cursor.fetchone()
    
    conn.close()
    return result

def get_instruction(drug_name: str, instructions_dict: Optional[Dict[str, str]] = None) -> str:
    """Get instruction for a drug (for backward compatibility)"""
    if instructions_dict:
        return instructions_dict.get(drug_name, "No instructions found for this medication.")
    
    result = search_drug(drug_name)
    if result:
        return result[1]
    return "No instructions found for this medication."
